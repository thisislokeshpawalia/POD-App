import os
import shutil
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def format_row(row, bg_hex, text_color_rgb=RGBColor(30, 41, 59), bold=False, font_size=9.5):
    for cell in row.cells:
        set_cell_background(cell, bg_hex)
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(font_size)
                run.font.color.rgb = text_color_rgb
                run.font.bold = bold

def create_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
    format_row(table.rows[0], "1E3A8A", text_color_rgb=RGBColor(255, 255, 255), bold=True, font_size=9.5)

    # Data rows
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
        bg = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
        format_row(row, bg, text_color_rgb=RGBColor(30, 41, 59), bold=False, font_size=9.0)

    # Apply column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph() # spacing after table
    return table

def add_heading_1(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Deep Navy #1E3A8A

def add_heading_2(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 125, 50) # Forest Green #2E7D32

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(30, 41, 59)
    return p

def add_bullet(doc, title, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run_b = p.add_run(title + ": ")
    run_b.font.name = 'Calibri'
    run_b.font.size = Pt(9.5)
    run_b.font.bold = True
    run_b.font.color.rgb = RGBColor(30, 41, 59)
    
    run_t = p.add_run(text)
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(9.5)
    run_t.font.color.rgb = RGBColor(30, 41, 59)
    return p

def generate_docx(filename):
    doc = Document()

    # Page Margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("Functional Requirements Specification (FRS)")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 58, 138)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(14)
    run_sub = p_sub.add_run("Subsidy Delivery Partner (Pod Delivery App) — Version 1.2")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(46, 125, 50)

    # 1. Document Control
    add_heading_1(doc, "1. Document Control")
    doc_control_headers = ["Attribute", "Details"]
    doc_control_data = [
        ["Application Name", "Subsidy Delivery Partner (Pod Delivery App)"],
        ["Document Name", "Functional Requirements Specification (FRS)"],
        ["Document Version", "1.2"],
        ["Prepared / Developed By", "Lokesh Pawalia and Sarthak Srivastava"],
        ["Prepared Date", "August 31, 2026"],
        ["Reviewed By", "Lokesh Pawalia & Sarthak Srivastava"],
        ["Approved By", "Lokesh Pawalia & Sarthak Srivastava"],
        ["Document Status", "Approved & Production Ready"],
        ["Deployment Environment", "Android APK Client & Railway Production Cloud"],
        ["Backend Architecture", "FastAPI (Python 3.11) + MongoDB Atlas (pody_db)"],
        ["Media CDN Storage", "Cloudinary (Images, Proof Videos, PDF Invoices)"]
    ]
    create_styled_table(doc, doc_control_headers, doc_control_data, col_widths=[2.5, 4.4])

    # 2. Purpose
    add_heading_1(doc, "2. Purpose")
    add_body(doc, "The purpose of this document is to define the end-to-end functional requirements for the Subsidy Delivery Partner (Pod Delivery) mobile application and its corresponding backend services. It details all system behaviors, expected user interactions, validation logic, biometric security gates, digital invoicing protocols, and cloud synchronization workflows.")
    add_body(doc, "Version 1.2 establishes a comprehensive, single source of truth for engineering stakeholders, delivery partners, and administrative officers, ensuring zero supply-chain leakage and foolproof fraud prevention in the distribution of subsidized agricultural materials to verified beneficiaries.")

    # 3. Scope
    add_heading_1(doc, "3. Scope")
    add_heading_2(doc, "In Scope (Version 1.2)")
    add_bullet(doc, "Delivery Partner Authentication", "Mobile-number-based authentication, SMS OTP verification, and JWT session handling.")
    add_bullet(doc, "Strict 8-Step Profile Onboarding", "Full Name character filtering, RFC email validation, 6-digit PIN code with automated Indian Postal API city/state lookup, Address, Vehicle Type, Indian Vehicle Registration number format validation, Aadhaar 12-digit UIDAI validation, and profile photo cloud upload.")
    add_bullet(doc, "Scoped Beneficiary Management", "Viewing partner-specific farmers filtered by All, Pending, and Delivered statuses, with on-the-fly farmer creation.")
    add_bullet(doc, "Geolocation Proximity Check", "Enforces a strict 600m geofence radius between the delivery agent's live GPS coordinates and the beneficiary's registered coordinates.")
    add_bullet(doc, "On-Device Facial Verification", "Local computer vision using Google ML Kit to match the live recipient's facial landmarks against the beneficiary reference photo.")
    add_bullet(doc, "Proof of Delivery (POD)", "Multi-angle photo capture (1 to 5 images), live video audit recording, and recipient delivery OTP.")
    add_bullet(doc, "Digital Invoicing & Native Inspection", "Automated PDF invoice generation, Cloudinary sync, and in-app viewing via Syncfusion Flutter PDF Viewer.")

    add_heading_2(doc, "Out of Scope")
    add_bullet(doc, "Public Beneficiary Portal", "Farmers do not use this app directly; they interact with delivery partners.")
    add_bullet(doc, "Payment Processing", "The platform strictly focuses on subsidy goods disbursal and cryptographic handover proof, not financial payment gateway processing.")

    # 4. Application Overview
    add_heading_1(doc, "4. Application Overview & Typical Workflow")
    add_body(doc, "The Subsidy Delivery Partner application is an Android mobile solution designed for delivery personnel to securely disburse subsidized agricultural goods to registered farmers. To prevent corruption and phantom deliveries, the app mandates five progressive verification checks.")
    add_body(doc, "Typical End-to-End Workflow:")
    add_body(doc, "1. Partner Login / Register with 10-digit mobile number → Verify OTP (1234).\n"
                    "2. If new user: Complete 8-Step Profile (Name, Email, PIN, Auto City/State, Address, Vehicle Type, Vehicle Number, Aadhaar, Photo) → Auto-redirected.\n"
                    "3. Dashboard (Customer List) → View assigned deliveries (Pending / Delivered).\n"
                    "4. Select Pending Farmer → View subsidy goods, contact details, and route coordinates.\n"
                    "5. Navigate to Location → Execute 600m Geofence Proximity Check.\n"
                    "6. Enter Beneficiary Handover OTP → Verify.\n"
                    "7. Capture Beneficiary Face Photo → On-device ML Kit Face Match.\n"
                    "8. Capture 1 to 5 Proof Photos + Optional Video Proof → Submit Order.\n"
                    "9. System marks order as Delivered, generates PDF invoice, uploads to Cloudinary, and allows instant native viewing via Syncfusion PDF Viewer.")

    # 5. User Roles
    add_heading_1(doc, "5. User Roles & Hierarchy")
    roles_headers = ["Role ID", "User Role", "Description", "System Permissions"]
    roles_data = [
        ["UR-01", "Delivery Partner", "Field agent transporting and disbursing subsidized equipment.", "Authenticate, complete profile, view assigned farmers, execute POD workflow, view invoices."],
        ["UR-02", "Beneficiary (Farmer)", "Registered farmer receiving allocated subsidy items.", "Provides OTP, physical presence for face match, receives subsidized items."],
        ["UR-03", "System Administrator", "Central supervisor monitoring system operations and audit logs.", "Access cloud database (MongoDB pody_db), review media assets on Cloudinary, monitor delivery logs."]
    ]
    create_styled_table(doc, roles_headers, roles_data, col_widths=[0.8, 1.4, 2.5, 2.2])

    # 6. Application Navigation
    add_heading_1(doc, "6. Application Navigation Architecture")
    add_body(doc, "Splash Screen\n"
                    "  └── Authentication Gate\n"
                    "        ├── [Unauthenticated] Login Screen / Register Flow (Mobile + OTP)\n"
                    "        ├── [Needs Profile] Profile Completion Screen (Strict 8-Step Form)\n"
                    "        └── [Authenticated] Permission Gate (Location + Camera)\n"
                    "              └── Dashboard (Customer List Screen)\n"
                    "                    ├── Filter Tabs (All / Pending / Delivered)\n"
                    "                    ├── Add Farmer Dialog (On-the-fly beneficiary registration)\n"
                    "                    ├── Customer Detail Screen (Items, Phone Call, Navigation)\n"
                    "                    │     └── Order Completion Screen (5 Verification Gates)\n"
                    "                    │           ├── Step 1: 600m Geolocation Verification\n"
                    "                    │           ├── Step 2: Recipient SMS OTP Verification\n"
                    "                    │           ├── Step 3: Local Face Verification (Google ML Kit)\n"
                    "                    │           ├── Step 4: Multi-Photo Proof (1-5 images)\n"
                    "                    │           ├── Step 5: Video Proof Capture\n"
                    "                    │           └── Order Finalization & Cloudinary Upload\n"
                    "                    └── In-App Native PDF Viewer Screen (Syncfusion PDF Viewer)")

    # 7. Functional Requirements Matrix
    add_heading_1(doc, "7. Functional Requirements Matrix")

    # 7.1 Auth Module
    add_heading_2(doc, "7.1 Authentication Module (AUTH)")
    auth_headers = ["Req ID", "Requirement", "Preconditions", "User Action", "Expected System Behavior", "Priority"]
    auth_data = [
        ["FR-AUTH-001", "Mobile OTP Request", "None", "Enter 10-digit mobile number", "Calls POST /auth/check-phone; if exists, sends OTP; if not, routes to register.", "P0"],
        ["FR-AUTH-002", "New Partner Registration", "Unregistered phone", "Enter Name, Phone, Aadhaar", "Validates input constraints and dispatches OTP via POST /auth/send-otp.", "P0"],
        ["FR-AUTH-003", "OTP Verification", "OTP dispatched", "Enter 4-digit OTP", "Verifies OTP via POST /auth/verify-otp. Generates JWT session token.", "P0"],
        ["FR-AUTH-004", "Persistent Session (Auto-Login)", "Valid stored token", "Launch application", "Validates token via GET /auth/me. Routes to dashboard or profile completion.", "P0"]
    ]
    create_styled_table(doc, auth_headers, auth_data, col_widths=[0.9, 1.4, 1.0, 1.1, 2.0, 0.5])

    # 7.2 Profile Module
    add_heading_2(doc, "7.2 Delivery Partner Profile Module (PROF)")
    prof_headers = ["Req ID", "Requirement", "Preconditions", "User Action", "Expected System Behavior", "Priority"]
    prof_data = [
        ["FR-PROF-001", "Mandatory Profile Completion", "Authenticated new partner", "Fill 8-step form and submit", "Enforces non-empty profile before allowing access to deliveries.", "P0"],
        ["FR-PROF-002", "Full Name Input Filter", "Profile form open", "Type in Name field", "Keyboard allows only [a-zA-Z\\s]. Blocks digits and special characters.", "P0"],
        ["FR-PROF-003", "Email Format Enforcement", "Profile form open", "Enter email address", "Validates RFC regex format ^[\\w-\\.]+@([\\w-]+\\.)+[\\w-]{2,4}$. Blocks whitespace.", "P0"],
        ["FR-PROF-004", "PIN Code Auto Geo-Lookup", "Profile form open", "Enter 6-digit PIN code", "Asynchronously calls Indian Postal API to resolve District (City) and State.", "P0"],
        ["FR-PROF-005", "City & State Auto-Population", "PIN resolved", "6-digit PIN entered", "Auto-fills City and State fields in read-only mode (with manual override option).", "P0"],
        ["FR-PROF-006", "Vehicle Number Format", "Profile form open", "Enter Vehicle Number", "Converts to uppercase; validates Indian registration pattern (e.g. DL 01 AB 1234 / BH).", "P0"],
        ["FR-PROF-007", "Aadhaar Number Validation", "Profile form open", "Enter 12-digit Aadhaar", "Validates exactly 12 numeric digits; rejects numbers starting with 0 or 1.", "P0"],
        ["FR-PROF-008", "Profile Photo Cloud Upload", "Profile form open", "Capture or select photo", "Uploads photo via multipart/form-data to PATCH /auth/profile and Cloudinary.", "P0"]
    ]
    create_styled_table(doc, prof_headers, prof_data, col_widths=[0.9, 1.4, 1.0, 1.1, 2.0, 0.5])

    # 7.3 Dashboard Module
    add_heading_2(doc, "7.3 Dashboard & Deliveries Module (DASH)")
    dash_headers = ["Req ID", "Requirement", "Preconditions", "User Action", "Expected System Behavior", "Priority"]
    dash_data = [
        ["FR-DASH-001", "Scoped Delivery Listing", "Authenticated partner", "Open Dashboard", "Queries GET /farmers?delivery_partner_id={id}. Shows categorized deliveries.", "P0"],
        ["FR-DASH-002", "Delivery Status Filter", "Deliveries loaded", "Tap filter chip", "Filters items dynamically by 'All', 'Pending', or 'Delivered'.", "P1"],
        ["FR-DASH-003", "On-the-Fly Beneficiary Add", "Partner logged in", "Tap '+ Add Farmer'", "Opens creation dialog; uploads photo to Cloudinary; saves to MongoDB.", "P1"],
        ["FR-DASH-004", "Direct Call & Navigation", "Farmer card open", "Tap Call or Map button", "Launches device dialer (tel:) or navigation maps (geo:) with target coordinates.", "P1"]
    ]
    create_styled_table(doc, dash_headers, dash_data, col_widths=[0.9, 1.4, 1.0, 1.1, 2.0, 0.5])

    # 7.4 Order Completion Module
    add_heading_2(doc, "7.4 Order Completion & Verification Protocol (ORD)")
    ord_headers = ["Req ID", "Requirement", "Preconditions", "User Action", "Expected System Behavior", "Priority"]
    ord_data = [
        ["FR-ORD-001", "600m Geofence Validation", "On completion view", "System checks GPS", "Calculates Haversine distance; blocks submission if distance > 600m.", "P0"],
        ["FR-ORD-002", "Customer Handover OTP", "Within 600m radius", "Enter 4-digit OTP", "Matches against beneficiary record. Blocks completion if incorrect.", "P0"],
        ["FR-ORD-003", "Local Biometric Face Match", "Within 600m radius", "Capture recipient face", "Google ML Kit compares live face with registered photo. Blocks on mismatch.", "P0"],
        ["FR-ORD-004", "Multi-Photo Proof Capture", "Face verified", "Capture 1-5 photos", "Stores images locally; uploads to Cloudinary upon submission.", "P0"],
        ["FR-ORD-005", "Live Video Audit Proof", "Photos captured", "Record video (<=30s)", "Encodes and streams video proof to Cloudinary POD-App/proofs.", "P1"],
        ["FR-ORD-006", "Automated PDF Invoice", "All verifications pass", "Tap Submit Delivery", "Generates PDF invoice, uploads to Cloudinary, marks status as 'delivered'.", "P0"],
        ["FR-ORD-007", "Native PDF Viewing", "Order delivered", "Tap 'View Invoice'", "Renders digital invoice in-app using Syncfusion PDF Viewer with pan & zoom.", "P0"]
    ]
    create_styled_table(doc, ord_headers, ord_data, col_widths=[0.9, 1.4, 1.0, 1.1, 2.0, 0.5])

    # 8. Screen Specifications
    add_heading_1(doc, "8. Screen-by-Screen Functional Specifications")
    add_heading_2(doc, "8.1 Login Screen")
    add_body(doc, "• Purpose: Authenticate existing delivery partners or route new operatives to registration.\n"
                    "• UI Elements: Clean brand illustration, 10-digit mobile number input, 'Login' button, 'Register' button.\n"
                    "• Validations: Mobile number must be exactly 10 digits; non-digit characters are blocked.")

    add_heading_2(doc, "8.2 Registration & OTP Screen")
    add_body(doc, "• Purpose: Onboard new operatives and verify contact credentials.\n"
                    "• UI Elements: Full Name (strict alphabetic filter), Mobile Number (+91 prefix), Aadhaar Number, 4-digit OTP inputs.\n"
                    "• Validations: Name must not contain digits/symbols; Aadhaar must be 12 digits; OTP is 4 digits (Testing OTP: 1234).")

    add_heading_2(doc, "8.3 Profile Completion Screen (Strict 8-Step Form)")
    add_body(doc, "• Purpose: Collect verified operational profile information before permitting delivery operations.\n"
                    "• UI Elements & Field Sequence:\n"
                    "   1. Full Name: Alphabetical only, max 50 chars, blocks numbers/symbols via FilteringTextInputFormatter.\n"
                    "   2. Email Address: Validated RFC format (^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\\.[a-zA-Z]{2,}$, spaces denied).\n"
                    "   3. PIN Code: Exactly 6 digits; auto-triggers background postal lookup with loading spinner.\n"
                    "   4. City & State: Read-only auto-populated fields from resolved PIN data (manual override toggle available).\n"
                    "   5. Address: Street/locality address, minimum 5 characters.\n"
                    "   6. Vehicle Type: Dropdown (Two Wheeler, Three Wheeler, Four Wheeler, Other).\n"
                    "   7. Vehicle Number: Auto-capitalized; validated against standard Indian vehicle registration formats.\n"
                    "   8. Aadhaar Number: Exactly 12 digits; cannot begin with 0 or 1 (UIDAI format).\n"
                    "   • Profile Photo: Mandatory circular avatar with camera/gallery picker; uploaded via multipart to Cloudinary.")

    add_heading_2(doc, "8.4 Dashboard (Customer List Screen)")
    add_body(doc, "• Purpose: Central operational hub displaying deliveries assigned to the logged-in delivery partner.\n"
                    "• UI Elements: Header greeting partner with live delivery tally ('X/Y Delivered'), search/filter chips (All, Pending, Delivered), Beneficiary Cards displaying name, village, contact, item count, and delivery status badge.\n"
                    "• Top Action: '+ Add Farmer' button opening beneficiary creation dialog.")

    add_heading_2(doc, "8.5 Order Completion Screen (5 Verification Gates)")
    add_body(doc, "• Purpose: Execute the physical handover verification workflow.\n"
                    "• UI Elements & Gating Logic:\n"
                    "   • Proximity Card: Real-time GPS distance calculation. Green if <= 600m; Red warning if > 600m.\n"
                    "   • Beneficiary OTP Input: Required 4-digit code.\n"
                    "   • Face Verification Button: Activates camera; Google ML Kit detects facial landmarks and compares against registered photo.\n"
                    "   • Item Proofs Grid: Supports capturing 1 to 5 delivery photos with thumbnail previews.\n"
                    "   • Video Proof Button: Records up to 30s video proof.\n"
                    "   • Submit Button: Finalizes order, syncs media to Cloudinary, updates MongoDB, and triggers PDF generation.")

    add_heading_2(doc, "8.6 Native PDF Viewer Screen")
    add_body(doc, "• Purpose: Inspect and download official delivery invoices directly inside the application.\n"
                    "• Technology: Syncfusion Flutter PDF Viewer (syncfusion_flutter_pdfviewer).\n"
                    "• Features: High-resolution vector rendering, pinch-to-zoom, page navigation, and download action.")

    # 9. API Specification Matrix
    add_heading_1(doc, "9. Cloud Backend API Specification Matrix")
    add_body(doc, "Base URL: https://pod-app-production-818a.up.railway.app | Protocol: HTTPS | Format: JSON / Multipart Form-Data")
    api_headers = ["Method", "Endpoint", "Purpose", "Payload Type", "Auth Required"]
    api_data = [
        ["POST", "/auth/check-phone", "Verify partner registration status", "JSON {phone}", "No"],
        ["POST", "/auth/send-otp", "Dispatch SMS OTP to mobile", "JSON {phone}", "No"],
        ["POST", "/auth/verify-otp", "Validate 4-digit OTP", "JSON {phone, otp}", "No"],
        ["POST", "/auth/login-or-register", "Generate session access token", "JSON {phone, password, name}", "No"],
        ["PATCH", "/auth/profile", "Update profile details & photo", "Multipart (fields + profile_image)", "Bearer JWT"],
        ["GET", "/auth/me", "Fetch authenticated partner profile", "None", "Bearer JWT"],
        ["GET", "/farmers", "List assigned deliveries", "Query params (?status_filter)", "Bearer JWT"],
        ["POST", "/farmers", "Register new farmer beneficiary", "Multipart (data JSON + photo file)", "Bearer JWT"],
        ["GET", "/farmers/{id}", "Retrieve single farmer record", "Path parameter", "Bearer JWT"],
        ["PATCH", "/farmers/{id}", "Update farmer details", "JSON payload", "Bearer JWT"],
        ["POST", "/farmers/{id}/upload_proof", "Submit POD (video, photos, invoice)", "Multipart (video, photos, invoice_pdf)", "Bearer JWT"],
        ["DELETE", "/farmers/{id}", "Delete farmer record", "Path parameter", "Bearer JWT"]
    ]
    create_styled_table(doc, api_headers, api_data, col_widths=[0.8, 1.8, 2.2, 1.4, 0.7])

    # 10. Security & Non-Functional Requirements
    add_heading_1(doc, "10. Security & Non-Functional Requirements")
    add_bullet(doc, "Cryptographic Authentication", "All private endpoints require Bearer JWT tokens signed with HS256 algorithm and a minimum 64-character secret key.")
    add_bullet(doc, "Password Security", "Passwords salted and hashed using Bcrypt before storage in MongoDB.")
    add_bullet(doc, "Biometric Privacy", "Facial recognition is performed locally on-device via Google ML Kit; raw biometric vectors are not transmitted across public networks.")
    add_bullet(doc, "Network Resilience", "API requests implement a 120-second streaming timeout to handle low-bandwidth rural connectivity during photo/video uploads.")
    add_bullet(doc, "Data Integrity", "All beneficiary records, partner credentials, and delivery audits are persisted in MongoDB Atlas (pody_db) with replica set fault tolerance.")

    # 11. Device Permissions
    add_heading_1(doc, "11. Device Permissions Specification")
    perm_headers = ["Permission", "Usage Purpose", "When Requested", "Enforcement Level"]
    perm_data = [
        ["ACCESS_FINE_LOCATION", "Calculates live distance to farmer coordinates for 600m geofence verification.", "App startup & Order Completion", "Strict Mandatory (Blocks submission if denied)"],
        ["CAMERA", "Captures profile photo, facial verification match, item proofs, and video.", "Profile setup & Order Completion", "Strict Mandatory (Blocks submission if denied)"],
        ["INTERNET", "Communicates with Railway API server, Indian Postal API, and Cloudinary.", "Continuous background operation", "Mandatory OS Level"],
        ["ACCESS_NETWORK_STATE", "Monitors cellular / Wi-Fi network connectivity state.", "Continuous background operation", "Mandatory OS Level"]
    ]
    create_styled_table(doc, perm_headers, perm_data, col_widths=[1.8, 2.6, 1.5, 1.0])

    # 12. Business Rules
    add_heading_1(doc, "12. Core Business Rules (BR)")
    add_bullet(doc, "BR-001 (Geofence Constraint)", "A delivery cannot be completed if the operative is greater than 600 meters from the beneficiary's registered coordinates.")
    add_bullet(doc, "BR-002 (Biometric Face Match)", "A delivery requires a verified facial match between the live recipient and the registered beneficiary reference photo.")
    add_bullet(doc, "BR-003 (Handover OTP)", "A delivery cannot be completed without entering the beneficiary's unique 4-digit handover OTP.")
    add_bullet(doc, "BR-004 (Photographic Evidence)", "A delivery requires at least 1 and up to 5 clear photographic proofs of the items being handed over.")
    add_bullet(doc, "BR-005 (Mandatory Profile Setup)", "Delivery partners cannot view or execute delivery orders until their 8-step profile is completed and verified.")
    add_bullet(doc, "BR-006 (Automatic Invoicing)", "Every verified delivery automatically compiles a PDF receipt, uploads it to Cloudinary, and links it to the MongoDB audit record.")

    # 13. Acceptance Criteria
    add_heading_1(doc, "13. Acceptance Criteria (Gherkin Scenarios)")
    add_heading_2(doc, "Scenario 1: Geofenced Delivery Check")
    add_body(doc, "Given the delivery partner is on the Order Completion screen\n"
                    "When their GPS location is calculated as 750 meters from the farmer's coordinates\n"
                    "Then the system shall display a red warning banner and disable the final submission action.")

    add_heading_2(doc, "Scenario 2: Automated PIN Code Resolution")
    add_body(doc, "Given the partner is completing their registration profile\n"
                    "When they enter a valid 6-digit PIN code '221108'\n"
                    "Then the system shall query the Indian Postal API, display a loading spinner, and automatically populate City as 'Varanasi' and State as 'Uttar Pradesh'.")

    add_heading_2(doc, "Scenario 3: Local Face Verification Match")
    add_body(doc, "Given the partner is completing a delivery for farmer 'Ramesh'\n"
                    "When the partner captures a live photo of the recipient\n"
                    "Then Google ML Kit shall compare landmark embeddings against the registered profile photo and grant clearance only if the match confidence threshold is satisfied.")

    # 14. Document Approvals
    add_heading_1(doc, "14. Document Approvals & Sign-Off")
    appr_headers = ["Project Role", "Name", "Sign-Off Status", "Approval Date"]
    appr_data = [
        ["Lead Developer & Architect", "Lokesh Pawalia", "Approved & Verified", "August 31, 2026"],
        ["Lead Developer & Backend Specialist", "Sarthak Srivastava", "Approved & Verified", "August 31, 2026"],
        ["Release Engineering", "Version 1.2 Production Release", "Released for Deployment", "August 31, 2026"]
    ]
    create_styled_table(doc, appr_headers, appr_data, col_widths=[2.0, 1.8, 1.6, 1.5])

    doc.save(filename)
    print(f"DOCX built successfully: {filename}")

if __name__ == "__main__":
    out_docx = "Subsidy_Delivery_Partner_FRS_v1.2.docx"
    generate_docx(out_docx)
    # Also copy to Downloads folder for immediate user convenience
    downloads_path = os.path.join(r"C:\Users\me\Downloads", out_docx)
    try:
        shutil.copyfile(out_docx, downloads_path)
        print(f"Copied to: {downloads_path}")
    except Exception as e:
        print(f"Could not copy to Downloads: {e}")
